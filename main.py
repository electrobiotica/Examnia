"""ExamGen AI – MVP backend con análisis por imagen, corrección inteligente, puntajes e idioma dinámico
Author: ChatGPT (OpenAI gpt-4o)
"""

import os
import json
import uuid
import base64
from typing import List, Any, Dict, Optional

from fastapi import FastAPI, HTTPException, UploadFile, File, Form
from fastapi.responses import FileResponse, HTMLResponse
from fastapi.staticfiles import StaticFiles
from fastapi.middleware.cors import CORSMiddleware
from pydantic import BaseModel
import openai
from docx import Document
from docx.shared import Pt

# === Config ===
openai.api_key = os.getenv("OPENAI_API_KEY")
if not openai.api_key:
    raise RuntimeError("❌ OPENAI_API_KEY no está configurada en el entorno.")
MODEL_GENERATE = "gpt-4o"
MODEL_GRADE = "gpt-4o"
MODEL_VISION = "gpt-4o"
FILES_DIR = "generated_files"
STATIC_DIR = "static"
UPLOADS_DIR = "uploads"

# Asegura carpetas
os.makedirs(FILES_DIR, exist_ok=True)
os.makedirs(STATIC_DIR, exist_ok=True)
os.makedirs(UPLOADS_DIR, exist_ok=True)

app = FastAPI(title="ExamGen AI", version="0.6.0")

app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

app.mount("/files", StaticFiles(directory=FILES_DIR), name="files")
app.mount("/static", StaticFiles(directory=STATIC_DIR), name="static")
app.mount("/lang", StaticFiles(directory=os.path.join(STATIC_DIR, "lang")), name="lang")

# === Schemas ===
class GenerateExamRequest(BaseModel):
    course: str
    topic: str
    objectives: str
    n_questions: int = 10
    q_type: str = "mcq"
    output_format: Optional[str] = "json"
    language: Optional[str] = "es"

class GenerateExamResponse(BaseModel):
    exam: List[Dict[str, Any]]
    download_url: Optional[str] = None

class GradeRequest(BaseModel):
    question: str
    rubric: str
    student_answer: str
    language: Optional[str] = "es"

class GradeResponse(BaseModel):
    score: float
    feedback: str

# === Prompts ===

def prompt_generate(req: GenerateExamRequest) -> str:
    return (
        f"You are an expert assessment designer. You must reply in {req.language}.\n"
        f"Create {req.n_questions} {req.q_type} questions for the following course.\n"
        f"Course: {req.course}\n"
        f"Topic/Unit: {req.topic}\n"
        f"Learning objectives: {req.objectives}\n"
        "For each question, include:\n"
        "- id\n"
        "- question\n"
        "- options (if type is mcq)\n"
        "- answer\n"
        "- rubric: use 3 levels (Excelente, Aceptable, Insuficiente) describing criteria for grading.\n"
        "Return only a JSON array."
    )


# === DOCX Generator ===
def build_docx(exam: List[Dict[str, Any]], meta: GenerateExamRequest) -> str:
    doc = Document()
    doc.add_heading(f"Examen – {meta.course}", level=0)
    doc.add_paragraph(f"Tema: {meta.topic}")
    doc.add_paragraph(f"Objetivos: {meta.objectives}")
    doc.add_paragraph("\n")
    for item in exam:
        q_run = doc.add_paragraph().add_run(f"{item['id']}. {item['question']}")
        q_run.font.size = Pt(12)
        if meta.q_type == "mcq":
            for letter, opt in zip("ABCD", item.get("options", [])):
                doc.add_paragraph(f"    {letter}) {opt}")
    doc.add_paragraph("\n\n---\nGenerado por ExamGen AI")
    filename = f"exam_{uuid.uuid4().hex}.docx"
    path = os.path.join(FILES_DIR, filename)
    doc.save(path)
    return filename

# === Endpoints ===
@app.get("/")
async def serve_index():
    index_path = os.path.join(STATIC_DIR, "index.html")
    if os.path.exists(index_path):
        with open(index_path, encoding="utf-8") as f:
            return HTMLResponse(content=f.read(), status_code=200)
    return HTMLResponse("<h1>Index.html no encontrado</h1>", status_code=404)

@app.post("/generate", response_model=GenerateExamResponse)
async def generate_exam(req: GenerateExamRequest):
    try:
        user_prompt = prompt_generate(req)
        response = openai.ChatCompletion.create(
            model=MODEL_GENERATE,
            messages=[
                {"role": "system", "content": f"You are an assessment generator that replies in {req.language}."},
                {"role": "user", "content": user_prompt},
            ],
            temperature=0.7,
            max_tokens=1500,
        )
        exam_json_str = response.choices[0].message.content.strip()

        # Extrae automáticamente el bloque de JSON si viene con ```json o similares
        matches = re.findall(r"""```(?:json)?\s*(\{.*?\}|\[.*?\])\s*```""", exam_json_str, re.DOTALL)
        if matches:
            exam_json_str = matches[0]
        exam = json.loads(exam_json_str)
        result: Dict[str, Any] = {"exam": exam}
        if req.output_format == "docx":
            filename = build_docx(exam, req)
            result["download_url"] = f"/files/{filename}"
        return result
    except Exception as e:
        print("❌ Error en /generate:", e)
        raise HTTPException(status_code=500, detail="Error generando el examen")

@app.post("/grade", response_model=GradeResponse)
async def grade_answer(req: GradeRequest):
    try:
        user_prompt = prompt_grade(req)
        response = openai.ChatCompletion.create(
            model=MODEL_GRADE,
            messages=[
                {"role": "system", "content": f"You are an exam corrector that responds in {req.language}."},
                {"role": "user", "content": user_prompt},
            ],
            temperature=0.2,
            max_tokens=300,
        )
        grade_json_str = response.choices[0].message.content.strip()
        grade_json = json.loads(grade_json_str)
        return grade_json
    except Exception as e:
        print("❌ Error en /grade:", e)
        raise HTTPException(status_code=500, detail="Error al calificar")

@app.post("/analyze-image")
async def analyze_image(file: UploadFile = File(...)):
    try:
        contents = await file.read()
        base64_image = base64.b64encode(contents).decode("utf-8")
        response = openai.ChatCompletion.create(
            model=MODEL_VISION,
            messages=[
                {"role": "user", "content": [
                    {"type": "text", "text": "You are an exam corrector. Analyze this photo, extract questions and answers, and prepare it for scoring. Answer in Spanish."},
                    {"type": "image_url", "image_url": {"url": f"data:image/jpeg;base64,{base64_image}"}}
                ]},
            ],
            max_tokens=1000,
        )
        result = response.choices[0].message.content.strip()
        return {"result": result}
    except Exception as e:
        print("❌ Error en /analyze-image:", e)
        raise HTTPException(status_code=500, detail="Error analizando la imagen del examen")



@app.post("/generate-key")
async def generate_answer_key(exam: List[Dict[str, Any]]):
    try:
        doc = Document()
        doc.add_heading("Hoja de respuestas", level=1)
        for item in exam:
            doc.add_paragraph(f"{item['id']}. {item['answer']}")
        filename = f"gabarito_{uuid.uuid4().hex}.docx"
        path = os.path.join(FILES_DIR, filename)
        doc.save(path)
        return {"download_url": f"/files/{filename}"}
    except Exception as e:
        print("❌ Error generando gabarito:", e)
        raise HTTPException(status_code=500, detail="Error al generar gabarito")

@app.get("/healthz")
async def health():
    return {"status": "ok"}


if __name__ == "__main__":
    import uvicorn
    uvicorn.run("main:app", host="0.0.0.0", port=8000, reload=True)
